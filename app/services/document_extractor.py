"""
Document Text Extractor Service
Extracts text from various document formats: PDF, DOCX, ODT, MSG, EML
"""
import os
from typing import Optional, Dict, Any
import fitz  # PyMuPDF
try:
    import pytesseract
    HAS_OCR = True
except ImportError:
    HAS_OCR = False
from PIL import Image
import io


class DocumentExtractor:
    """Unified document text extraction service"""
    
    SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.doc', '.odt', '.msg', '.eml'}
    
    def __init__(self):
        self._check_dependencies()

    def _init_tesseract(self):
        """Find tesseract executable on Windows"""
        if not HAS_OCR:
            return False

        if os.name == 'nt':
            # Default Windows paths
            possible_paths = [
                r"C:\Program Files\Tesseract-OCR\tesseract.exe",
                r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
                os.path.join(os.getcwd(), "tesseract-ocr", "tesseract.exe"),
                os.path.join(os.getcwd(), "tesseract-setup", "tesseract.exe") # If extracted here
            ]
            for p in possible_paths:
                if os.path.exists(p):
                    pytesseract.pytesseract.tesseract_cmd = p
                    print(f"DEBUG: Tesseract found at {p}")
                    return True
            print("WARNING: Tesseract-OCR not found in default paths.")
        return False
    
    def _check_dependencies(self):
        """Check and note available extraction libraries"""
        self.has_pymupdf = False
        self.has_docx = False
        self.has_odf = False
        self.has_msg = False
        
        try:
            import fitz  # PyMuPDF
            self.has_pymupdf = True
        except ImportError:
            pass
        
        try:
            import docx
            self.has_docx = True
        except ImportError:
            pass
        
        try:
            from odf import text, teletype
            from odf.opendocument import load as load_odf
            self.has_odf = True
        except ImportError:
            pass
        
        try:
            import extract_msg
            self.has_msg = True
        except ImportError:
            pass
    
    def extract(self, filepath: str) -> Dict[str, Any]:
        """
        Extract text and metadata from a document
        
        Returns:
            dict with keys: text, metadata, success, error
        """
        if not os.path.exists(filepath):
            return {'success': False, 'error': 'Datei nicht gefunden', 'text': '', 'metadata': {}}
        
        ext = os.path.splitext(filepath)[1].lower()
        
        try:
            if ext == '.pdf':
                return self._extract_pdf(filepath)
            elif ext in ('.docx', '.doc'):
                return self._extract_docx(filepath)
            elif ext == '.odt':
                return self._extract_odt(filepath)
            elif ext == '.msg':
                return self._extract_msg(filepath)
            elif ext == '.eml':
                return self._extract_eml(filepath)
            else:
                return {'success': False, 'error': f'Format {ext} nicht unterstützt', 'text': '', 'metadata': {}}
        except Exception as e:
            return {'success': False, 'error': str(e), 'text': '', 'metadata': {}}
    
    def _extract_pdf(self, filepath: str) -> Dict[str, Any]:
        """Extract text from PDF using PyMuPDF"""
        if not self.has_pymupdf:
            return {'success': False, 'error': 'PyMuPDF nicht installiert', 'text': '', 'metadata': {}}
        
        import fitz
        
        text_parts = []
        metadata = {}
        
        doc = fitz.open(filepath)
        metadata = {
            'title': doc.metadata.get('title', ''),
            'author': doc.metadata.get('author', ''),
            'subject': doc.metadata.get('subject', ''),
            'creator': doc.metadata.get('creator', ''),
            'producer': doc.metadata.get('producer', ''),
            'page_count': doc.page_count,
            'creation_date': doc.metadata.get('creationDate', ''),
        }
        
        self._init_tesseract()
        
        for page_num in range(doc.page_count):
            page = doc[page_num]
            page_text = page.get_text().strip()
            
            # OCR Fallback: If text is empty or very short (< 50 chars), try OCR
            if len(page_text) < 50:
                try:
                    # Render page to image
                    pix = page.get_pixmap(dpi=300)
                    img_data = pix.tobytes("png")
                    image = Image.open(io.BytesIO(img_data))
                    
                    # Run OCR (German assumed + English)
                    if HAS_OCR:
                        ocr_text = pytesseract.image_to_string(image, lang='deu+eng')
                        
                        if len(ocr_text.strip()) > len(page_text):
                            page_text = f"--- OCR EXTRACTED (Page {page_num+1}) ---\n{ocr_text}"
                            print(f"DEBUG: OCR extracted {len(ocr_text)} chars on page {page_num+1}.")
                    else:
                        print("DEBUG: OCR skipped (pytesseract not installed)")
                except Exception as e:
                    print(f"WARNING: OCR failed for page {page_num+1}: {e}")
            
            text_parts.append(page_text)
        
        doc.close()
        
        return {
            'success': True,
            'text': '\n\n'.join(text_parts),
            'metadata': metadata,
            'error': None
        }
    
    def _extract_docx(self, filepath: str) -> Dict[str, Any]:
        """Extract text from DOCX using python-docx"""
        if not self.has_docx:
            return {'success': False, 'error': 'python-docx nicht installiert. Installieren mit: pip install python-docx', 'text': '', 'metadata': {}}
        
        from docx import Document
        
        doc = Document(filepath)
        
        # Extract text from paragraphs
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)
        
        # Get metadata from core properties
        metadata = {}
        try:
            core_props = doc.core_properties
            metadata = {
                'title': core_props.title or '',
                'author': core_props.author or '',
                'subject': core_props.subject or '',
                'created': str(core_props.created) if core_props.created else '',
                'modified': str(core_props.modified) if core_props.modified else '',
            }
        except:
            pass
        
        return {
            'success': True,
            'text': '\n\n'.join(text_parts),
            'metadata': metadata,
            'error': None
        }
    
    def _extract_odt(self, filepath: str) -> Dict[str, Any]:
        """Extract text from ODT using odfpy"""
        if not self.has_odf:
            return {'success': False, 'error': 'odfpy nicht installiert. Installieren mit: pip install odfpy', 'text': '', 'metadata': {}}
        
        from odf.opendocument import load as load_odf
        from odf import text as odf_text
        from odf import teletype
        
        doc = load_odf(filepath)
        
        # Extract all text elements
        text_parts = []
        for element in doc.getElementsByType(odf_text.P):
            content = teletype.extractText(element)
            if content.strip():
                text_parts.append(content)
        
        # Get metadata
        metadata = {}
        try:
            meta = doc.meta
            if meta:
                # ODT uses Dublin Core metadata
                for child in meta.childNodes:
                    if hasattr(child, 'qname'):
                        tag = child.qname[1] if isinstance(child.qname, tuple) else str(child.qname)
                        if hasattr(child, 'firstChild') and child.firstChild:
                            metadata[tag] = str(child.firstChild)
        except:
            pass
        
        return {
            'success': True,
            'text': '\n\n'.join(text_parts),
            'metadata': metadata,
            'error': None
        }
    
    def _extract_msg(self, filepath: str) -> Dict[str, Any]:
        """Extract text from Outlook MSG files"""
        if not self.has_msg:
            return {'success': False, 'error': 'extract-msg nicht installiert', 'text': '', 'metadata': {}}
        
        import extract_msg
        
        msg = extract_msg.Message(filepath)
        
        text_parts = []
        
        # Add subject
        if msg.subject:
            text_parts.append(f"Betreff: {msg.subject}")
        
        # Add sender
        if msg.sender:
            text_parts.append(f"Von: {msg.sender}")
        
        # Add date
        if msg.date:
            text_parts.append(f"Datum: {msg.date}")
        
        # Add body
        if msg.body:
            text_parts.append("")
            text_parts.append(msg.body)
        
        metadata = {
            'subject': msg.subject or '',
            'sender': msg.sender or '',
            'date': str(msg.date) if msg.date else '',
            'to': msg.to or '',
            'cc': msg.cc or '',
        }
        
        msg.close()
        
        return {
            'success': True,
            'text': '\n'.join(text_parts),
            'metadata': metadata,
            'error': None
        }
    
    def _extract_eml(self, filepath: str) -> Dict[str, Any]:
        """Extract text from EML email files"""
        import email
        from email import policy
        from email.parser import BytesParser
        
        with open(filepath, 'rb') as f:
            msg = BytesParser(policy=policy.default).parse(f)
        
        text_parts = []
        
        # Add headers
        subject = msg.get('Subject', '')
        sender = msg.get('From', '')
        date = msg.get('Date', '')
        
        if subject:
            text_parts.append(f"Betreff: {subject}")
        if sender:
            text_parts.append(f"Von: {sender}")
        if date:
            text_parts.append(f"Datum: {date}")
        
        # Get body
        body = ""
        if msg.is_multipart():
            for part in msg.walk():
                content_type = part.get_content_type()
                if content_type == 'text/plain':
                    try:
                        body = part.get_content()
                        break
                    except:
                        pass
        else:
            try:
                body = msg.get_content()
            except:
                body = msg.get_payload(decode=True).decode('utf-8', errors='ignore')
        
        if body:
            text_parts.append("")
            text_parts.append(body)
        
        metadata = {
            'subject': subject,
            'sender': sender,
            'date': date,
            'to': msg.get('To', ''),
            'cc': msg.get('Cc', ''),
        }
        
        return {
            'success': True,
            'text': '\n'.join(text_parts),
            'metadata': metadata,
            'error': None
        }
    
    def get_supported_extensions(self) -> set:
        """Return set of supported file extensions"""
        return self.SUPPORTED_EXTENSIONS
    
    def is_supported(self, filename: str) -> bool:
        """Check if file extension is supported"""
        ext = os.path.splitext(filename)[1].lower()
        return ext in self.SUPPORTED_EXTENSIONS
